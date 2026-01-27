import io
import re
from datetime import date, datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

import ms_graph
import sp_folder_graph as spg


# ==============================
# CONFIG
# ==============================
TEMPLATE_PATH = "Incident Report Template_blank (1).docx"

CITY_CODES = {
    "Davao City": "DVO",
    "Quezon City": "QZN",
}

STANDARD_IMAGE_WIDTH_IN = 5.5

SHAREPOINT_SITE_URL = st.secrets.get("sharepoint", {}).get("site_url", "")
INCIDENT_REPORTS_ROOT_PATH = st.secrets.get("sharepoint", {}).get(
    "incident_reports_root_path",
    "Ground Station Operations/Installations, Maintenance and Repair/Incident Reports",
)

SCOPES = ms_graph.DEFAULT_SCOPES_WRITE


# ==============================
# DOCX HELPERS
# ==============================
def _clear_table_rows_except_header(table, header_rows=1):
    while len(table.rows) > header_rows:
        tbl = table._tbl
        tr = table.rows[-1]._tr
        tbl.remove(tr)


def _set_2col_table_value(table, label, value):
    for row in table.rows:
        if row.cells and row.cells[0].text.strip() == label.strip():
            row.cells[1].text = "" if value is None else str(value)
            return


def _set_paragraph_after_heading(doc, heading_text, new_text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text.strip():
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].text = new_text or ""
            return


def _insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _append_figures_after_heading(doc, heading_text, files, captions, figure_start, section_label):
    if not files:
        return figure_start

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text.strip():
            anchor = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else p
            fig_no = figure_start

            for idx, f in enumerate(files):
                img_p = _insert_paragraph_after(anchor)
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(io.BytesIO(f.getvalue()), width=Inches(STANDARD_IMAGE_WIDTH_IN))

                caption_text = ""
                if captions and idx < len(captions):
                    caption_text = (captions[idx] or "").strip()
                if not caption_text:
                    caption_text = f.name.rsplit(".", 1)[0]

                cap_p = _insert_paragraph_after(img_p)
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_p.add_run(f"Figure {fig_no}. {section_label} – {caption_text}")
                cap_run.italic = True

                anchor = cap_p
                fig_no += 1

            return fig_no

    return figure_start


def _fill_sequence_table(table, df):
    _clear_table_rows_except_header(table, header_rows=0)
    for _, r in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r.get("Date", ""))
        cells[1].text = str(r.get("Time", ""))          # ✅ FIXED HERE
        cells[2].text = str(r.get("Category", ""))
        cells[3].text = str(r.get("Message", ""))


def _fill_actions_table(table, df):
    _clear_table_rows_except_header(table, header_rows=1)
    for _, r in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r.get("Date", ""))
        cells[1].text = str(r.get("Time", ""))
        cells[2].text = str(r.get("Performed by", ""))
        cells[3].text = str(r.get("Action", ""))
        cells[4].text = str(r.get("Result", ""))


def generate_docx(data):
    doc = Document(TEMPLATE_PATH)
    t0, t1, t2, t3 = doc.tables[0], doc.tables[1], doc.tables[2], doc.tables[3]

    _set_2col_table_value(t0, "Reported by", data["reported_by"])
    _set_2col_table_value(t0, "Position", data["position"])
    _set_2col_table_value(t0, "Date of Report", data["date_of_report"])
    _set_2col_table_value(t0, "Incident No.", data["full_incident_no"])

    _set_2col_table_value(t1, "Date (YYYY-MM-DD)", data["incident_date"])
    _set_2col_table_value(t1, "Time", data["incident_time"])
    _set_2col_table_value(t1, "Location", data["location"])
    _set_2col_table_value(t1, "Current Status", data["current_status"])

    _set_paragraph_after_heading(doc, "Nature of Incident", data["nature"])
    _set_paragraph_after_heading(doc, "Damages Incurred (if any)", data["damages"])
    _set_paragraph_after_heading(doc, "Investigation and Analysis", data["investigation"])
    _set_paragraph_after_heading(doc, "Conclusion and Recommendations", data["conclusion"])

    _fill_sequence_table(t2, data["sequence_df"])
    _fill_actions_table(t3, data["actions_df"])

    fig = 1
    fig = _append_figures_after_heading(doc, "Sequence of Events", data["sequence_images"], data["sequence_captions"], fig, "Sequence of Events")
    fig = _append_figures_after_heading(doc, "Damages Incurred (if any)", data["damages_images"], data["damages_captions"], fig, "Damages Incurred")
    fig = _append_figures_after_heading(doc, "Investigation and Analysis", data["investigation_images"], data["investigation_captions"], fig, "Investigation and Analysis")
    fig = _append_figures_after_heading(doc, "Conclusion and Recommendations", data["conclusion_images"], data["conclusion_captions"], fig, "Conclusion and Recommendations")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# ==============================
# PARSE EXISTING DOCX
# ==============================
def _get_2col_table_value(table, label):
    for row in table.rows:
        if row.cells and row.cells[0].text.strip() == label.strip():
            return row.cells[1].text.strip()
    return ""


def _get_paragraph_after_heading(doc, heading_text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text.strip():
            if i + 1 < len(doc.paragraphs):
                return doc.paragraphs[i + 1].text.strip()
            return ""
    return ""


def _table_to_sequence_df(table):
    rows = []
    for r in table.rows:
        cells = [c.text.strip() for c in r.cells]
        if len(cells) >= 4:
            rows.append({"Date": cells[0], "Time": cells[1], "Category": cells[2], "Message": cells[3]})
    df = pd.DataFrame(rows)
    if df.empty:
        df = pd.DataFrame([{"Date": "", "Time": "", "Category": "", "Message": ""}])
    return df


def _table_to_actions_df(table):
    rows = []
    for idx, r in enumerate(table.rows):
        cells = [c.text.strip() for c in r.cells]
        if len(cells) >= 5:
            if idx == 0 and ("Performed" in cells[2] or "Action" in cells[3] or "Result" in cells[4]):
                continue
            rows.append({"Date": cells[0], "Time": cells[1], "Performed by": cells[2], "Action": cells[3], "Result": cells[4]})
    df = pd.DataFrame(rows)
    if df.empty:
        df = pd.DataFrame([{"Date": "", "Time": "", "Performed by": "", "Action": "", "Result": ""}])
    return df


def parse_existing_ir_docx(docx_bytes: bytes) -> dict:
    doc = Document(io.BytesIO(docx_bytes))
    t0, t1, t2, t3 = doc.tables[0], doc.tables[1], doc.tables[2], doc.tables[3]

    return {
        "reported_by": _get_2col_table_value(t0, "Reported by"),
        "position": _get_2col_table_value(t0, "Position"),
        "date_of_report": _get_2col_table_value(t0, "Date of Report"),
        "full_incident_no": _get_2col_table_value(t0, "Incident No."),
        "incident_date": _get_2col_table_value(t1, "Date (YYYY-MM-DD)"),
        "incident_time": _get_2col_table_value(t1, "Time"),
        "location": _get_2col_table_value(t1, "Location"),
        "current_status": _get_2col_table_value(t1, "Current Status"),
        "nature": _get_paragraph_after_heading(doc, "Nature of Incident"),
        "damages": _get_paragraph_after_heading(doc, "Damages Incurred (if any)"),
        "investigation": _get_paragraph_after_heading(doc, "Investigation and Analysis"),
        "conclusion": _get_paragraph_after_heading(doc, "Conclusion and Recommendations"),
        "sequence_df": _table_to_sequence_df(t2),
        "actions_df": _table_to_actions_df(t3),
    }


# ==============================
# UI HELPERS
# ==============================
def captions_editor(files, key):
    if not files:
        return []
    df = pd.DataFrame({"File": [f.name for f in files], "Caption": ["" for _ in files]})
    edited = st.data_editor(df, key=key, num_rows="fixed", use_container_width=True)
    return edited["Caption"].tolist()


def normalize_serial(serial_raw: str) -> str:
    s = (serial_raw or "").strip()
    if not s:
        return ""
    if not re.fullmatch(r"\d{1,4}", s):
        return ""
    return s.zfill(4)


def _must_have_token():
    ms_graph.login_ui(scopes=SCOPES)
    token = ms_graph.get_access_token()
    if not token:
        st.stop()
    return token


def _ensure_defaults():
    defaults = {
        "reported_by": "",
        "position": "",
        "date_of_report": date.today().strftime("%Y-%m-%d"),
        "incident_date": date.today().strftime("%Y-%m-%d"),
        "incident_time": datetime.now().strftime("%H:%M:%S"),
        "location": "",
        "current_status": "Resolved",
        "nature": "",
        "damages": "None",
        "investigation": "",
        "conclusion": "",
        "seq_df": pd.DataFrame([{"Date": "", "Time": "", "Category": "", "Message": ""}]),
        "actions_df": pd.DataFrame([{"Date": "", "Time": "", "Performed by": "", "Action": "", "Result": ""}]),
        "serial_raw": "",
        "loaded_update_target": None,
        "loaded_full_incident_no": "",
    }
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)


def _df_valid(df: object) -> bool:
    return isinstance(df, pd.DataFrame) and (not df.empty) and (len(df.columns) > 0)


# ==============================
# APP START
# ==============================
st.set_page_config(page_title="Incident Report Generator", layout="wide")
st.title("Incident Report Generator")

token = _must_have_token()

if not SHAREPOINT_SITE_URL:
    st.error("Missing sharepoint.site_url in Streamlit secrets.")
    st.stop()

if "sp_site_id" not in st.session_state or "sp_drive_id" not in st.session_state:
    with st.spinner("Resolving SharePoint site/drive..."):
        st.session_state["sp_site_id"] = spg.resolve_site_id(token, SHAREPOINT_SITE_URL)
        st.session_state["sp_drive_id"] = spg.get_default_drive_id(token, st.session_state["sp_site_id"])

drive_id = st.session_state["sp_drive_id"]
this_year = str(datetime.now().year)

_ensure_defaults()

mode = st.radio("Mode", ["Create New", "Update Existing"], horizontal=True)

# ==============================
# UPDATE EXISTING SELECTOR
# ==============================
if mode == "Update Existing":
    st.subheader("Select existing Incident Report to update")

    u_year = st.selectbox("Year", [this_year, str(int(this_year) - 1)], index=0, key="u_year")
    u_city = st.selectbox("Ground Station Location", list(CITY_CODES.keys()), key="u_city")

    base_path = f"{INCIDENT_REPORTS_ROOT_PATH}/{u_year}/{u_city}"

    if st.button("Refresh folders/files", key="u_refresh"):
        for k in ["u_folders", "u_files", "u_files_folder"]:
            st.session_state.pop(k, None)

    if "u_folders" not in st.session_state:
        try:
            st.session_state["u_folders"] = spg.list_incident_folders(token, drive_id, base_path)
        except Exception as e:
            st.error(f"Cannot list incident folders: {e}")
            st.session_state["u_folders"] = []

    folders = st.session_state.get("u_folders", [])
    folder_names = [f["name"] for f in folders]

    u_folder_name = st.selectbox("Incident Folder (Incident No.)", ["-- select --"] + folder_names, key="u_folder")

    if u_folder_name != "-- select --":
        folder_meta = next((x for x in folders if x["name"] == u_folder_name), None)
        folder_id = folder_meta["id"]

        if "u_files" not in st.session_state or st.session_state.get("u_files_folder") != folder_id:
            try:
                st.session_state["u_files"] = spg.list_files(token, drive_id, folder_id)
                st.session_state["u_files_folder"] = folder_id
            except Exception as e:
                st.error(f"Cannot list files: {e}")
                st.session_state["u_files"] = []

        files = st.session_state.get("u_files", [])
        docx_files = [f for f in files if f["name"].lower().endswith(".docx")]
        docx_names = [f["name"] for f in docx_files]

        u_docx = st.selectbox("DOCX file", ["-- select --"] + docx_names, key="u_docx")

        if u_docx != "-- select --":
            fmeta = next((x for x in docx_files if x["name"] == u_docx), None)

            if st.button("Load into form", key="u_load"):
                try:
                    b = spg.download_file_bytes(token, drive_id, fmeta["id"])
                    parsed = parse_existing_ir_docx(b)

                    st.session_state["reported_by"] = parsed.get("reported_by", "")
                    st.session_state["position"] = parsed.get("position", "")
                    st.session_state["date_of_report"] = parsed.get("date_of_report", date.today().strftime("%Y-%m-%d"))

                    st.session_state["incident_date"] = parsed.get("incident_date", date.today().strftime("%Y-%m-%d"))
                    st.session_state["incident_time"] = parsed.get("incident_time", datetime.now().strftime("%H:%M:%S"))
                    st.session_state["location"] = parsed.get("location", u_city)
                    st.session_state["current_status"] = parsed.get("current_status", "Resolved") or "Resolved"

                    st.session_state["nature"] = parsed.get("nature", "")
                    st.session_state["damages"] = parsed.get("damages", "None") or "None"
                    st.session_state["investigation"] = parsed.get("investigation", "")
                    st.session_state["conclusion"] = parsed.get("conclusion", "")

                    seq_df = parsed.get("sequence_df")
                    act_df = parsed.get("actions_df")
                    if _df_valid(seq_df):
                        st.session_state["seq_df"] = seq_df
                    if _df_valid(act_df):
                        st.session_state["actions_df"] = act_df

                    st.session_state["loaded_update_target"] = {
                        "year": u_year,
                        "city": u_city,
                        "folder_name": u_folder_name,
                        "folder_id": folder_id,
                        "docx_name": u_docx,
                        "docx_id": fmeta["id"],
                    }
                    st.session_state["loaded_full_incident_no"] = parsed.get("full_incident_no", u_folder_name)

                    st.success("Loaded. Scroll down, edit details, then click Generate Report.")
                except Exception as e:
                    st.error(f"Load failed: {e}")

st.divider()

# ==============================
# MAIN FORM
# ==============================
loaded = st.session_state.get("loaded_update_target")

if mode == "Create New":
    year = st.selectbox("Year folder", [this_year, str(int(this_year) - 1)], index=0, key="main_year")
    city = st.selectbox("Ground Station Location", list(CITY_CODES.keys()), key="main_city")
    site_code = CITY_CODES[city]

    serial_raw = st.text_input("Incident serial (000#)", value=st.session_state.get("serial_raw", ""), key="serial_raw")
    serial = normalize_serial(serial_raw)

    full_incident_no = f"SMCOD-IR-GS-{site_code}-{year}-{serial}" if serial else ""
    st.text_input("Full Incident No. (auto)", value=full_incident_no, disabled=True)

else:
    if not loaded:
        st.warning("Load an existing DOCX above first.")
        st.stop()

    full_incident_no = st.session_state.get("loaded_full_incident_no") or loaded.get("folder_name", "")
    st.text_input("Incident No.", value=full_incident_no, disabled=True)

with st.form("ir_form"):
    c1, c2 = st.columns(2)
    with c1:
        reported_by = st.text_input("Reported by", key="reported_by")
        position = st.text_input("Position", key="position")
        date_of_report = st.text_input("Date of Report (YYYY-MM-DD)", key="date_of_report")
    with c2:
        incident_date = st.text_input("Incident Date (YYYY-MM-DD)", key="incident_date")
        incident_time = st.text_input("Incident Time", key="incident_time")
        location = st.text_input("Location", key="location")
        current_status = st.selectbox(
            "Current Status",
            ["Resolved", "Ongoing", "Monitoring", "Open"],
            index=["Resolved", "Ongoing", "Monitoring", "Open"].index(st.session_state.get("current_status", "Resolved")),
            key="current_status",
        )

    nature = st.text_area("Nature of Incident", height=120, key="nature")

    st.subheader("Sequence of Events")
    seq_df = st.data_editor(
        st.session_state.get("seq_df"),
        num_rows="dynamic",
        use_container_width=True,
        key="seq_editor",
    )
    seq_imgs = st.file_uploader("Sequence Photos (optional)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    seq_caps = captions_editor(seq_imgs or [], "seq_caps")

    damages = st.text_area("Damages Incurred", key="damages")
    dmg_imgs = st.file_uploader("Damage Photos (optional)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    dmg_caps = captions_editor(dmg_imgs or [], "dmg_caps")

    investigation = st.text_area("Investigation and Analysis", height=120, key="investigation")
    inv_imgs = st.file_uploader("Investigation Photos (optional)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    inv_caps = captions_editor(inv_imgs or [], "inv_caps")

    conclusion = st.text_area("Conclusion and Recommendations", height=120, key="conclusion")
    con_imgs = st.file_uploader("Conclusion Photos (optional)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    con_caps = captions_editor(con_imgs or [], "con_caps")

    st.subheader("Response and Actions Taken")
    actions_df = st.data_editor(
        st.session_state.get("actions_df"),
        num_rows="dynamic",
        use_container_width=True,
        key="actions_editor",
    )

    submit = st.form_submit_button("Generate Report")

if submit:
    if mode == "Create New":
        serial = normalize_serial(st.session_state.get("serial_raw", ""))
        if not serial:
            st.error("Enter a valid incident serial (numbers only up to 4 digits). Example: 0001 or 1.")
            st.stop()

    data = {
        "reported_by": reported_by,
        "position": position,
        "date_of_report": date_of_report,
        "full_incident_no": full_incident_no,
        "incident_date": incident_date,
        "incident_time": incident_time,
        "location": location,
        "current_status": current_status,
        "nature": nature,
        "damages": damages,
        "investigation": investigation,
        "conclusion": conclusion,
        "sequence_df": seq_df,
        "actions_df": actions_df,
        "sequence_images": seq_imgs or [],
        "damages_images": dmg_imgs or [],
        "investigation_images": inv_imgs or [],
        "conclusion_images": con_imgs or [],
        "sequence_captions": seq_caps or [],
        "damages_captions": dmg_caps or [],
        "investigation_captions": inv_caps or [],
        "conclusion_captions": con_caps or [],
    }

    docx_bytes = generate_docx(data)

    try:
        with st.spinner("Uploading DOCX to SharePoint..."):
            if mode == "Update Existing":
                target_folder_id = loaded["folder_id"]
                filename = loaded.get("docx_name") or f"{full_incident_no}.docx"
            else:
                is_dup = spg.check_duplicate_ir(
                    token,
                    drive_id,
                    INCIDENT_REPORTS_ROOT_PATH,
                    st.session_state["main_year"],
                    st.session_state["main_city"],
                    full_incident_no,
                )
                if is_dup:
                    st.error("Duplicate found: this Incident No folder already exists. Use a new serial.")
                    st.stop()

                incident_folder = spg.ensure_path(
                    token,
                    drive_id,
                    INCIDENT_REPORTS_ROOT_PATH,
                    parts=[st.session_state["main_year"], st.session_state["main_city"], full_incident_no],
                )
                target_folder_id = incident_folder["id"]
                filename = f"{full_incident_no}.docx"

            spg.upload_file_to_folder(
                token,
                drive_id,
                folder_item_id=target_folder_id,
                filename=filename,
                content_bytes=docx_bytes,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        st.success("Report generated and uploaded.")
    except Exception as e:
        st.error(f"Upload failed: {e}")

    st.download_button(
        "Download Incident Report (DOCX)",
        data=docx_bytes,
        file_name=f"{full_incident_no}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
